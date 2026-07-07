from src.tools.extract_requirements import MAX_FILE_SIZE_BYTES, extract_requirements


def test_missing_file_returns_error(tmp_path):
    result = extract_requirements(str(tmp_path / "does_not_exist.txt"))
    assert "error" in result
    assert "not found" in result["error"].lower()


def test_unsupported_extension_returns_error(tmp_path):
    path = tmp_path / "requirements.exe"
    path.write_bytes(b"junk")

    result = extract_requirements(str(path))

    assert "error" in result
    assert "unsupported" in result["error"].lower()


def test_file_too_large_returns_error(tmp_path):
    path = tmp_path / "big.txt"
    path.write_bytes(b"x" * (MAX_FILE_SIZE_BYTES + 1))

    result = extract_requirements(str(path))

    assert "error" in result
    assert "too large" in result["error"].lower()


def test_txt_extraction(tmp_path):
    path = tmp_path / "req.txt"
    path.write_text("The system shall allow login with valid credentials.", encoding="utf-8")

    result = extract_requirements(str(path))

    assert result["file_type"] == "txt"
    assert "shall allow login" in result["raw_text"]
    assert result["char_count"] == len(result["raw_text"])


def test_csv_extraction_returns_rows(tmp_path):
    path = tmp_path / "req.csv"
    path.write_text(
        "id,requirement,priority\nREQ-1,Login with valid creds,high\nREQ-2,Reject invalid password,medium\n",
        encoding="utf-8",
    )

    result = extract_requirements(str(path))

    assert result["file_type"] == "csv"
    assert result["row_count"] == 2
    assert result["rows"][0] == {"id": "REQ-1", "requirement": "Login with valid creds", "priority": "high"}


def test_docx_extraction_groups_by_heading(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_heading("Login", level=1)
    doc.add_paragraph("User must be able to log in with valid credentials.")
    doc.add_heading("Logout", level=1)
    doc.add_paragraph("User must be able to log out from any page.")
    path = tmp_path / "req.docx"
    doc.save(str(path))

    result = extract_requirements(str(path))

    assert result["file_type"] == "docx"
    headings = [s["heading"] for s in result["sections"]]
    assert headings == ["Login", "Logout"]
    assert "log in with valid credentials" in result["sections"][0]["text"]


def test_xlsx_extraction_returns_sheet_rows(tmp_path):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Requirements"
    ws.append(["id", "requirement"])
    ws.append(["REQ-1", "Login with valid creds"])
    path = tmp_path / "req.xlsx"
    wb.save(str(path))

    result = extract_requirements(str(path))

    assert result["file_type"] == "xlsx"
    assert result["sheets"]["Requirements"] == [{"id": "REQ-1", "requirement": "Login with valid creds"}]
    assert result["row_count"] == 1


def test_pdf_extraction_reports_page_count(tmp_path):
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    path = tmp_path / "req.pdf"
    with path.open("wb") as f:
        writer.write(f)

    result = extract_requirements(str(path))

    assert result["file_type"] == "pdf"
    assert result["page_count"] == 1
